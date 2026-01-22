"""
Word Exporter Module for DPMPTSP Reporting System

This module generates Word (.docx) reports with charts, tables, and narratives
using python-docx library.
"""

import io
from pathlib import Path
from typing import Dict, Optional
from datetime import datetime

# Try to import Word generation library
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordExporter:
    """
    Exports reports to Word (.docx) format using python-docx.
    
    Features:
    - Professional layout with header
    - Embedded charts as images
    - Styled tables with data
    - Narrative sections with formatting
    """
    
    # Color scheme (RGB tuples)
    COLORS = {
        'primary': RGBColor(30, 58, 95) if DOCX_AVAILABLE else None,
        'secondary': RGBColor(61, 126, 166) if DOCX_AVAILABLE else None,
        'white': RGBColor(255, 255, 255) if DOCX_AVAILABLE else None,
    }
    
    def __init__(self, logo_path: Optional[Path] = None):
        self.logo_path = logo_path
    
    def is_available(self) -> bool:
        """Check if Word export is available."""
        return DOCX_AVAILABLE
    
    def export_report(
        self,
        report,
        stats: Dict,
        narratives,
        charts: Dict[str, bytes],
        output_path: Optional[Path] = None
    ) -> bytes:
        """
        Export a complete report to Word document.
        
        Args:
            report: PeriodReport object
            stats: Summary statistics dictionary
            narratives: Narrative object with all text sections
            charts: Dictionary of chart names to PNG bytes
            output_path: Optional path to save Word file
            
        Returns:
            Word document as bytes
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Please install it with: pip install python-docx")
        
        # Create document
        doc = Document()
        
        # Set up document styles
        self._setup_styles(doc)
        
        # Set page margins
        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        
        # 1. Cover Page
        self._add_cover_page(doc, report)
        doc.add_page_break()
        
        # 2. Table of Contents
        self._add_table_of_contents(doc)
        doc.add_page_break()
        
        # Add metrics summary table
        self._add_metrics_table(doc, stats)
        
        # Section 1: Pendahuluan
        self._add_section_title(doc, "Pendahuluan")
        self._add_paragraph(doc, narratives.pendahuluan)
        
        # Section 2: Rekapitulasi NIB
        self._add_section_title(doc, "1.1 Rekapitulasi Data NIB")
        
        if 'monthly' in charts:
            self._add_chart_image(doc, charts['monthly'])
        
        self._add_paragraph(doc, narratives.rekapitulasi_nib)
        
        # Section 3: Per Kabupaten/Kota
        self._add_section_title(doc, "1.2 Rekapitulasi per Kabupaten/Kota")
        
        if 'kab_kota' in charts:
            self._add_chart_image(doc, charts['kab_kota'])
        
        self._add_paragraph(doc, narratives.rekapitulasi_kab_kota)
        
        # Add data table
        self._add_subsection_title(doc, "Tabel Data per Kabupaten/Kota")
        self._add_data_table(doc, stats)
        
        # Section 4: Status PM
        doc.add_page_break()
        self._add_section_title(doc, "1.3 Status Penanaman Modal")
        
        if 'pm' in charts:
            self._add_chart_image(doc, charts['pm'], width=4)
        
        self._add_paragraph(doc, narratives.status_pm)
        
        # Section 5: Pelaku Usaha
        self._add_section_title(doc, "1.4 Kategori Pelaku Usaha")
        
        if 'pelaku' in charts:
            self._add_chart_image(doc, charts['pelaku'], width=4)
        
        self._add_paragraph(doc, narratives.pelaku_usaha)
        
        # Section 6: Sektor & Risiko (if data available)
        sektor_risiko = stats.get('sektor_risiko', {})
        if sektor_risiko:
            self._add_section_title(doc, "1.5 Perizinan Berdasarkan Risiko dan Sektor")
            
            if 'risk' in charts:
                self._add_chart_image(doc, charts['risk'], width=4)
            
            if 'sector' in charts:
                self._add_chart_image(doc, charts['sector'], width=5)
            
            # Generate sektor risiko narrative
            sektor_narrative = self._generate_sektor_risiko_narrative(sektor_risiko)
            self._add_paragraph(doc, sektor_narrative)
        
        # ============== SECTION 2: INVESTASI/PROYEK ==============
        doc.add_page_break()
        self._add_section_title(doc, "2. Rekapitulasi Data Investasi dan Proyek")
        
        # 2.1 PMA/PMDN Proyek
        if 'proyek_pm' in charts:
            self._add_subsection_title(doc, "2.1 Distribusi Proyek PMA/PMDN")
            self._add_chart_image(doc, charts['proyek_pm'], width=4)
        
        # 2.3 Skala Usaha
        if 'skala_usaha' in charts:
            self._add_subsection_title(doc, "2.3 Proyek Berdasarkan Skala Usaha")
            self._add_chart_image(doc, charts['skala_usaha'], width=5)
        
        # ============== SECTION 3: PERIZINAN BERUSAHA ==============
        doc.add_page_break()
        self._add_section_title(doc, "3. Perizinan Berusaha Berbasis Risiko")
        
        # 3.1 Kab/Kota PB
        if 'pb_kab_kota' in charts:
            self._add_subsection_title(doc, "3.1 Perizinan per Kabupaten/Kota")
            self._add_chart_image(doc, charts['pb_kab_kota'], width=5)
        
        # 3.2 Status PM PB
        if 'pb_pm' in charts:
            self._add_subsection_title(doc, "3.2 Perizinan Berdasarkan Status PM")
            self._add_chart_image(doc, charts['pb_pm'], width=4)
        
        # 3.3 Risk Level PB
        if 'pb_risk' in charts:
            self._add_subsection_title(doc, "3.3 Perizinan Berdasarkan Tingkat Risiko")
            self._add_chart_image(doc, charts['pb_risk'], width=5)
        
        # 3.4 Sector PB
        if 'pb_sector' in charts:
            self._add_subsection_title(doc, "3.4 Top 10 Sektor Perizinan")
            self._add_chart_image(doc, charts['pb_sector'], width=5)
        
        # Section 7: Kesimpulan
        doc.add_page_break()
        self._add_section_title(doc, "Kesimpulan")
        self._add_paragraph(doc, narratives.kesimpulan)
        
        # 8. Closing Page
        doc.add_page_break()
        self._add_closing_page(doc)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        docx_bytes = buffer.getvalue()
        buffer.close()
        
        # Save to file if path provided
        if output_path:
            with open(output_path, 'wb') as f:
                f.write(docx_bytes)
        
        return docx_bytes
    
    def _add_cover_page(self, doc, report):
        """Add cover page with logo, title, and metadata."""
        # Add logo if available
        if self.logo_path and self.logo_path.exists():
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(self.logo_path), width=Inches(4.5))
                doc.add_paragraph()  # Spacer
            except Exception:
                pass
        
        # Add spacer
        for _ in range(3):
            doc.add_paragraph()
        
        # Main title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("LAPORAN REKAPITULASI\nDATA NOMOR INDUK BERUSAHA (NIB)")
        run.bold = True
        run.font.size = Pt(24)
        run.font.color.rgb = self.COLORS['primary']
        
        doc.add_paragraph()
        
        # Period subtitle
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f"PERIODE {report.period_name} TAHUN {report.year}")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = self.COLORS['secondary']
        
        # Add spacer
        for _ in range(5):
            doc.add_paragraph()
        
        # Metadata table
        metadata = [
            ("Tim Pengolah Data", "DPMPTSP Provinsi Lampung"),
            ("Sumber Data", "OSS-RBA (Online Single Submission)"),
            ("Tanggal Penarikan Data", datetime.now().strftime("%d %B %Y")),
        ]
        
        for label, value in metadata:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run1 = p.add_run(f"{label}: ")
            run1.bold = True
            run1.font.size = Pt(11)
            run2 = p.add_run(value)
            run2.font.size = Pt(11)
    
    def _add_table_of_contents(self, doc):
        """Add table of contents page."""
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("DAFTAR ISI")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = self.COLORS['primary']
        
        doc.add_paragraph()
        
        toc_items = [
            "1. Nomor Induk Berusaha (NIB)",
            "   1.1 Rekapitulasi Data NIB",
            "   1.2 Rekapitulasi per Kabupaten/Kota",
            "   1.3 Status Penanaman Modal",
            "   1.4 Kategori Pelaku Usaha",
            "   1.5 Perizinan Berdasarkan Risiko dan Sektor",
            "2. Kesimpulan",
        ]
        
        for item in toc_items:
            p = doc.add_paragraph()
            run = p.add_run(item)
            run.font.size = Pt(12)
    
    def _add_closing_page(self, doc):
        """Add closing page with thank you message and contact info."""
        # Spacer
        for _ in range(6):
            doc.add_paragraph()
        
        # Thank you message
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TERIMA KASIH")
        run.bold = True
        run.font.size = Pt(36)
        run.font.color.rgb = self.COLORS['primary']
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add logo if available
        if self.logo_path and self.logo_path.exists():
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(self.logo_path), width=Inches(3))
            except Exception:
                pass
        
        doc.add_paragraph()
        
        # Contact info
        contact_info = [
            "DPMPTSP Provinsi Lampung",
            "Jl. Wolter Monginsidi No. 69, Bandar Lampung",
            "Telepon: (0721) 123456",
            "Website: https://dpmptsp.lampungprov.go.id",
            "Email: dpmptsp@lampungprov.go.id",
        ]
        
        for line in contact_info:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(10)
    
    def _generate_sektor_risiko_narrative(self, sektor_risiko_data: dict) -> str:
        """Generate narrative for Sektor & Risiko section."""
        total_risiko = (
            sektor_risiko_data.get('risiko_rendah', 0) +
            sektor_risiko_data.get('risiko_menengah_rendah', 0) +
            sektor_risiko_data.get('risiko_menengah_tinggi', 0) +
            sektor_risiko_data.get('risiko_tinggi', 0)
        )
        
        if total_risiko == 0:
            return "Data perizinan berdasarkan risiko belum tersedia."
        
        risk_levels = {
            'Rendah': sektor_risiko_data.get('risiko_rendah', 0),
            'Menengah Rendah': sektor_risiko_data.get('risiko_menengah_rendah', 0),
            'Menengah Tinggi': sektor_risiko_data.get('risiko_menengah_tinggi', 0),
            'Tinggi': sektor_risiko_data.get('risiko_tinggi', 0),
        }
        dominant_risk = max(risk_levels, key=risk_levels.get)
        dominant_pct = (risk_levels[dominant_risk] / total_risiko * 100) if total_risiko > 0 else 0
        
        sectors = {
            'Perindustrian': sektor_risiko_data.get('sektor_perindustrian', 0),
            'Kelautan & Perikanan': sektor_risiko_data.get('sektor_kelautan', 0),
            'Pertanian': sektor_risiko_data.get('sektor_pertanian', 0),
        }
        sectors_filtered = {k: v for k, v in sectors.items() if v > 0}
        
        if sectors_filtered:
            dominant_sector = max(sectors_filtered, key=sectors_filtered.get)
            sector_text = f"Sektor {dominant_sector} menempati posisi tertinggi."
        else:
            sector_text = ""
        
        narrative = f'Berdasarkan tingkat risiko, perizinan dengan kategori "{dominant_risk}" mendominasi dengan {risk_levels[dominant_risk]:,} perizinan ({dominant_pct:.1f}%). {sector_text}'
        
        return narrative.replace(',', '.')
    
    def _setup_styles(self, doc):
        """Setup document styles."""
        # Modify Normal style
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)
    
    def _add_header(self, doc, report):
        """Add header with logo and title."""
        # Add logo if available
        if self.logo_path and self.logo_path.exists():
            try:
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(str(self.logo_path), width=Inches(5))
            except Exception:
                pass
        
        # Add title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = title.add_run("LAPORAN REKAPITULASI DATA NIB")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = self.COLORS['primary']
        
        title.add_run("\n")
        
        run2 = title.add_run(f"PERIODE {report.period_name} TAHUN {report.year}")
        run2.bold = True
        run2.font.size = Pt(14)
        run2.font.color.rgb = self.COLORS['primary']
        
        doc.add_paragraph()  # Spacer
    
    def _add_metrics_table(self, doc, stats: Dict):
        """Add metrics summary table."""
        total_nib = stats.get('total_nib', 0)
        pm_dist = stats.get('pm_distribution', {})
        pelaku = stats.get('pelaku_usaha_distribution', {})
        
        table = doc.add_table(rows=2, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        headers = ['Total NIB', 'PMDN', 'PMA', 'UMK']
        values = [
            f"{total_nib:,}".replace(',', '.'),
            f"{pm_dist.get('PMDN', 0):,}".replace(',', '.'),
            f"{pm_dist.get('PMA', 0):,}".replace(',', '.'),
            f"{pelaku.get('UMK', 0):,}".replace(',', '.')
        ]
        
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._set_cell_shading(cell, "1e3a5f")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = self.COLORS['white']
        
        for i, value in enumerate(values):
            cell = table.rows[1].cells[i]
            cell.text = value
            self._set_cell_shading(cell, "f8f9fa")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(12)
        
        doc.add_paragraph()  # Spacer
    
    def _set_cell_shading(self, cell, color: str):
        """Set cell background color."""
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _add_section_title(self, doc, text: str):
        """Add a section title."""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = self.COLORS['primary']
        
        # Add bottom border
        paragraph.paragraph_format.space_after = Pt(10)
    
    def _add_subsection_title(self, doc, text: str):
        """Add a subsection title."""
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = self.COLORS['secondary']
    
    def _add_paragraph(self, doc, text: str):
        """Add a body paragraph."""
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(text)
        run.font.size = Pt(11)
        paragraph.paragraph_format.space_after = Pt(12)
    
    def _add_chart_image(self, doc, chart_bytes: bytes, width: float = 6):
        """Add a chart image to the document."""
        img_buffer = io.BytesIO(chart_bytes)
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(img_buffer, width=Inches(width))
        doc.add_paragraph()  # Spacer
    
    def _add_data_table(self, doc, stats: Dict):
        """Add data table for top 5 kabupaten/kota."""
        top_5 = stats.get('top_5_locations', [])
        if not top_5:
            return
        
        total_nib = stats.get('total_nib', 1)
        
        table = doc.add_table(rows=len(top_5) + 1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        headers = ['No', 'Kabupaten/Kota', 'Total NIB', 'Persentase']
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            self._set_cell_shading(cell, "1e3a5f")
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = self.COLORS['white']
                    run.font.size = Pt(10)
        
        # Data rows
        for row_idx, loc in enumerate(top_5, 1):
            pct = (loc['Total'] / total_nib * 100) if total_nib > 0 else 0
            row_data = [
                str(row_idx),
                loc['Kabupaten/Kota'],
                f"{loc['Total']:,}".replace(',', '.'),
                f"{pct:.1f}%"
            ]
            
            for col_idx, value in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                cell.text = value
                # Alternate row colors
                if row_idx % 2 == 0:
                    self._set_cell_shading(cell, "f8f9fa")
                for paragraph in cell.paragraphs:
                    if col_idx == 1:  # Kabupaten/Kota left-aligned
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    else:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(10)
        
        doc.add_paragraph()  # Spacer
